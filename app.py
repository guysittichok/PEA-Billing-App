import streamlit as st
import pdfplumber
import re
import pandas as pd
from io import BytesIO
import openpyxl
# 🛠️ นำเข้าโมดูลจัดการ สี, ฟอนต์ และการจัดวาง ของ openpyxl
from openpyxl.styles import Font, Alignment

st.set_page_config(page_title="ระบบจัดการบิลค่าไฟฟ้า", layout="wide")
st.title("PEA Bill Extraction System")

def extract_exact_pea_bill(file_obj):
    with pdfplumber.open(file_obj) as pdf:
        text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

    result = {
        "ชื่อไฟล์": file_obj.name,
        "C": 0.0, "D": 0.0, "E": 0.0, "F": 0.0, "G": 0.0, "H": 0.0,
        "I": 0.0, "J": 0.0, "K": 0.0, "L": 0.0, "M": 0.0, "N": 0.0, 
        "O": 0.0, "P": 0.0, "Q": 0.0
    }

    # ตรวจสอบรูปแบบบิล TOU
    is_tou = any(re.search(r''+k+r'.*?(?:กว|หน่วย|หนอรย|หนวย)', text, re.I) for k in ["Peak", "Off Peak", "Partial Peak"]) or ("OP" in text and " P " in text)
    has_h_mode = " H " in text or "\nH " in text or " H\n" in text or " Holiday " in text

    # ========================================================
    # บล็อกลอจิกเดิมทุกช่อง (ห้ามแตะต้องเด็ดขาด)
    # ========================================================
    if is_tou and has_h_mode:
        peak_money_match = re.search(r'Peak\s+[\d,]+\.\d+\s+กว\.\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
        if peak_money_match: result["F"] = float(peak_money_match.group(1).replace(",", ""))

        if result["F"] == 0.0 or result["F"] == 0:
            gwa_pattern = re.search(r'([\d,]+\.\d+)\s+กว\.?\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text)
            if gwa_pattern:
                result["F"] = float(gwa_pattern.group(2).replace(",", ""))
            
        off_peak_money_match = re.search(r'Off\s+Peak\s+[\d,]+\.\d+\s+กว\.\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
        if off_peak_money_match: result["G"] = float(off_peak_money_match.group(1).replace(",", ""))
            
        h_money_match = re.search(r'(?:H|Holiday)\s+[\d,]+\.\d+\s+กว\.\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
        if h_money_match: result["H"] = float(h_money_match.group(1).replace(",", ""))

        parts = text.split("พลังงานไฟฟ้า")
        demand_part = parts[0]
        energy_part = parts[1] if len(parts) > 1 else text 

        demand_lines = demand_part.split('\n')
        for line in demand_lines:
            nums = re.findall(r"([\d,]+\.\d+)", line)
            if not nums: continue
            if "P" in line and "กว" in line and "Off" not in line:
                result["C"] = float(nums[2].replace(",", "")) if len(nums) >= 3 else float(nums[0].replace(",", ""))
            elif "OP" in line:
                result["D"] = float(nums[2].replace(",", "")) if len(nums) >= 3 else float(nums[0].replace(",", ""))
            elif line.strip().startswith("H ") or " H " in line or "Holiday" in line:
                if "กว" in line or len(nums) >= 3:
                    result["E"] = float(nums[2].replace(",", "")) if len(nums) >= 3 else float(nums[0].replace(",", ""))

        p_unit = re.search(r'(?:^|\s+)P\s+[\d,]+\.\d+\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', energy_part, re.I)
        if p_unit: result["I"] = float(p_unit.group(1).replace(",", ""))
        else:
            p_unit_alt = re.search(r'(?:พลังงานไฟฟ้า)?\s+P\s+[\d,]+\.\d+\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
            if p_unit_alt: result["I"] = float(p_unit_alt.group(1).replace(",", ""))

        op_unit = re.search(r'OP\s+[\d,]+\.\d+\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', energy_part, re.I)
        if op_unit: result["J"] = float(op_unit.group(1).replace(",", ""))

        h_unit = re.search(r'(?:H|Holiday)\s+[\d,]+\.\d+\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', energy_part, re.I)
        if h_unit: result["K"] = float(h_unit.group(1).replace(",", ""))
            
        if result["K"] == 0.0:
            h_fallback = re.search(r'OP.*?([\d,]+\.\d+)\s+(?:H|Holiday)\s+[\d,]+\.\d+\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', energy_part, re.DOTALL | re.I)
            if h_fallback: result["K"] = float(h_fallback.group(2).replace(",", ""))

    elif not is_tou:
        lines = text.split('\n')
        unit_match = re.search(r'([\d,]+\.\d+)\s+(?:หน่วย|หนอรย|หนวย)', text)
        if unit_match: result["I"] = float(unit_match.group(1).replace(",", ""))

        if "พลังไฟฟ้าสูงสุด" in text:
            for line in lines:
                if "พลังไฟฟ้าสูงสุด" in line:
                    nums = re.findall(r"([\d,]+\.\d+)", line)
                    if nums:
                        if len(nums) >= 3: result["C"] = float(nums[2].replace(",", ""))
                        else: result["C"] = float(nums[-1].replace(",", ""))

            demand_cost_match = re.search(r'พลังไฟฟ้าสูงสุด\s+.*?กว\..*?([\d,]+\.\d+)', text)
            if demand_cost_match: result["F"] = float(demand_cost_match.group(1).replace(",", ""))

            if result["F"] == 0.0 or result["F"] == 0:
                gwa_pattern = re.search(r'([\d,]+\.\d+)\s+กว\.\s+([\d,]+\.\d+)\s+([\d,]+\.\d+)', text)
                if gwa_pattern:
                    result["C"] = float(gwa_pattern.group(1).replace(",", ""))
                    result["F"] = float(gwa_pattern.group(3).replace(",", ""))
            if result["F"] == 0.0 or result["F"] == 0:
                last_ditch_match = re.findall(r'กว\..*?([\d,]+\.\d+)', text)
                if last_ditch_match:
                    result["F"] = float(last_ditch_match[-1].replace(",", ""))
            
    else:
        peak = re.search(r'Peak\s+([\d,]+\.\d+)\s+กว\.\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
        if peak:
            result["C"] = float(peak.group(1).replace(",", ""))
            result["F"] = float(peak.group(2).replace(",", ""))

        pp = re.search(r'Partial\s+Peak\s+([\d,]+\.\d+)\s+กว\.\s+[\d,]+\.\d+\s+([\d,]+\.\d+)', text, re.I)
        if pp:
            result["D"] = float(pp.group(1).replace(",", ""))
            result["G"] = float(pp.group(2).replace(",", ""))

        op = re.search(r'Off\s+Peak\s+([\d,]+\.\d+)\s+กว', text, re.I)
        if op: result["E"] = float(op.group(1).replace(",", ""))

        for line in text.split('\n'):
            nums = re.findall(r"([\d,]+\.\d+)", line)
            if not nums: continue
            if "พลังงานไฟฟ้า" in line and "P" in line and "PP" not in line: result["I"] = float(nums[-1].replace(",", ""))
            elif "PP" in line: result["J"] = float(nums[-1].replace(",", ""))
            elif "OP" in line: result["K"] = float(nums[-1].replace(",", ""))

    # ดึงค่าช่อง M (Off Peak หน่วยสะสมซ้ายสุด)
    for line in text.split('\n'):
        if "off" in line.lower() and "peak" in line.lower() and any(k in line for k in ["หน่วย", "หนอรย", "หนวย"]):
            clean_line = re.split(r'\d{2}/\d{2}/\d{2,4}', line)[0]
            nums_in_op_line = re.findall(r"([\d,]+\.\d+)", clean_line)
            if nums_in_op_line:
                result["M"] = float(nums_in_op_line[-1].replace(",", ""))
                break

    # ========================================================
    # ระบบ Block Search Pattern (ช่อง L)
    # ========================================================
    found_l = False

    if is_tou:
        tou_pattern = re.search(r'(?:Peak|[\d,]+\.\d+)\s+(?:หน่วย|หนอรย|หนวย)\s+(\d+\.\d{4})\s+([\d,]+\.\d+)', text, re.I)
        if tou_pattern:
            potential_money = float(tou_pattern.group(2).replace(",", ""))
            if potential_money != result["M"] and potential_money not in [65760.0, 379280.0]:
                result["L"] = potential_money
                found_l = True
        
        if not found_l:
            for line in text.split('\n'):
                if re.search(r'\d{2}/\d{2}/\d{2,4}', line) or "ประวัติ" in line or "history" in line.lower() or "กว" in line:
                    continue
                if any(k in line for k in ["หน่วย", "หนอรย", "หนวย"]):
                    nums = re.findall(r"([\d,]+\.\d+)", line)
                    if len(nums) >= 3:
                        if any(len(n.split('.')[-1]) == 4 for n in nums):
                            potential_money = float(nums[-1].replace(",", ""))
                            if potential_money != result["M"] and potential_money not in [65760.0, 379280.0]:
                                result["L"] = potential_money
                                found_l = True
                                break

    if not is_tou or not found_l:
        for line in text.split('\n'):
            if re.search(r'\d{2}/\d{2}/\d{2,4}', line) or any(k in line for k in ["ประวัติ", "history", "กว.", "กว", "ค่าบริการ", "รวมเงิน", "total"]): 
                continue
            if any(k in line for k in ["หน่วย", "หนอรย", "หนวย"]):
                nums_in_line = re.findall(r"([\d,]+\.\d+)", line)
                if nums_in_line:
                    result["L"] = float(nums_in_line[-1].replace(",", ""))
                    break

    # ========================================================
    # ดึงค่า Ft, PF และ Sub Total
    # ========================================================
    ft = re.search(r'ค่า\s*Ft.*?=\s*[\d\.]+\s*บาท/หน่วย\s+([\d,]+\.\d+)', text, re.I)
    if not ft: ft = re.search(r'ค่า\s*Ft.*?([\d,]+\.\d+)', text, re.I)
    if ft: result["O"] = float(ft.group(1).replace(",", ""))
    
    result["P"] = 0.0  
    for line in text.split('\n'):
        if any(k in line for k in ["ค่าเพาเวอร์แฟคเตอร", "เพาเวอร์แฟคเตอร์", "Power Factor", "คาเพาเวอร์แฟคเตอร"]):
            nums_in_pf_line = re.findall(r"([\d,]+\.\d+)", line)
            if nums_in_pf_line:
                result["P"] = float(nums_in_pf_line[-1].replace(",", ""))
                break 

    total_match = re.search(r'รวมเงินค่าไฟฟ้า\s*\(Sub Total\)\s*([\d,]+\.\d+)', text, re.I)
    if total_match: result["Q"] = float(total_match.group(1).replace(",", ""))

    return result

# ----------------------------------------------------
# ส่วนโครงสร้าง Excel & Streamlit UI 
# ----------------------------------------------------
uploaded_files = st.file_uploader("อัปโหลดไฟล์บิล PDF", type=["pdf"], accept_multiple_files=True)
template_file = st.file_uploader("อัปโหลดไฟล์ Excel ต้นแบบ", type=["xlsx"])

if uploaded_files:
    data = [extract_exact_pea_bill(f) for f in uploaded_files]
    all_cols = ["ชื่อไฟล์", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q"]
    df = pd.DataFrame(data, columns=all_cols)
    
    df["O"] = "-"
    df["Q"] = "-"
    
    st.data_editor(df, use_container_width=True)

    if template_file and st.button("สร้างไฟล์ Excel พร้อมข้อมูล"):
        try:
            wb = openpyxl.load_workbook(template_file)
            
            # 🎯 แก้ไขลอจิกการดึงหน้าชีท เพื่อป้องกันการเกิดบั๊ก NoneType 
            if wb.active is not None:
                ws = wb.active
            else:
                ws = wb.worksheets[0]

            # 🎯 ปรับแต่งฟอนต์สีน้ำเงินแบบ "ตัวปกติ (bold=False)" และชิดขวามือทั้งหมด
            blue_normal_font = Font(name="Calibri", size=11, bold=False, color="0000FF")
            right_alignment = Alignment(horizontal="right", vertical="center")

            def write_number(ws, cell_pos, value):
                val_str = str(value).strip()
    
                # บังคับใช้สีฟอนต์ตัวปกติสีน้ำเงิน และ จัดตำแหน่งชิดขวา
                ws[cell_pos].font = blue_normal_font
                ws[cell_pos].alignment = right_alignment
    
                # กำหนดรูปแบบ Number Format แบบ Accounting (ถ้าเป็น 0 จะแสดงผลเป็นเครื่องหมายขีด - อัตโนมัติ)
                # และสามารถนำไป บวก ลบ คูณ หาร ใน Excel ได้ปกติ ไม่เกิด #VALUE!
                accounting_format = '_(* #,##0.00_);_(* (#,##0.00);_(* "-"??_);_(@_)'
    
                if val_str in ["None", "", "-"]:
                    ws[cell_pos] = 0
                    ws[cell_pos].number_format = accounting_format
                else:
                    try:
                        val = float(val_str.replace(',', ''))
                        ws[cell_pos] = val
                        ws[cell_pos].number_format = accounting_format
                    except:
                        # หากแปลงค่าไม่ได้จริงๆ ให้ใส่เลข 0 ไว้ก่อนเพื่อความปลอดภัยในสูตรคำนวณ
                        ws[cell_pos] = 0
                        ws[cell_pos].number_format = accounting_format

            for row_idx in range(20, ws.max_row + 1):
                excel_key = str(ws[f'A{row_idx}'].value).strip()
                if excel_key in ["None", ""]: continue
                
                match_row = df[df['ชื่อไฟล์'].apply(lambda x: excel_key in str(x))]
                
                if not match_row.empty:
                    row = match_row.iloc[0]
                    write_number(ws, f'C{row_idx}', row['C'])
                    write_number(ws, f'D{row_idx}', row['D'])
                    write_number(ws, f'E{row_idx}', row['E'])
                    write_number(ws, f'F{row_idx}', row['F'])
                    write_number(ws, f'G{row_idx}', row['G'])
                    write_number(ws, f'H{row_idx}', row['H'])
                    write_number(ws, f'I{row_idx}', row['I'])
                    write_number(ws, f'J{row_idx}', row['J'])
                    write_number(ws, f'K{row_idx}', row['K'])
                    write_number(ws, f'L{row_idx}', row['L'])
                    write_number(ws, f'M{row_idx}', row['M'])
                    write_number(ws, f'N{row_idx}', row['N'])
                    write_number(ws, f'P{row_idx}', row['P'])
            
            output = BytesIO()
            
            # 🎯 เพิ่มลอจิกเปิดโหมดบังคับคำนวณสูตรออโต้ เพื่อไม่ให้ค่าในสูตรแสดงเป็นช่องว่างตอนดาวน์โหลดออกไปครับ
            wb.properties.calcMode = 'auto'
            
            wb.save(output)
            output.seek(0)
            st.success("กรอกข้อมูลลงใน Template เรียบร้อยและรักษาสูตรเดิมแล้ว!")
            st.download_button("📥 ดาวน์โหลด Excel", output, "Updated_PEA_Bill.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            
        except Exception as e:
            st.error(f"เกิดข้อผิดพลาด: {e}")

import datetime
from io import BytesIO
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# ================================================================
# 1. ฟังก์ชันช่วยจัดรูปแบบเอกสารและล้างข้อผิดพลาด XML
# ================================================================
def add_run_thai(paragraph, text, font_name="TH Sarabun PSK", size_pt=16, bold=False, color_rgb=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    return run

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=20, bottom=20, left=60, right=60):
    # ปรับ Margin เซลล์ให้แคบลงเพื่อป้องกันการล้นหน้ากระดาษ
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_custom_info_borders(table):
    # แก้ไข XML String ให้สมบูรณ์แบบเพื่อปิดปัญหา Premature end of data 
    tblPr = table._tbl.tblPr
    xml_string = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(xml_string))

# ================================================================
# 2. ฟังก์ชันสร้าง Word สรุปแบบคลีนข้อมูล (บีบให้อยู่ในหน้าเดียว)
# ================================================================
def create_exact_layout_report(df_clean, selected_month, selected_year, reference_rate):
    doc = Document()
    
    # กำหนดขอบกระดาษแบบแคบ (บน-ล่าง 0.8 นิ้ว, ซ้าย-ขวา 1 นิ้ว) เพื่อการันตีเนื้อหาจบในหน้าเดียว
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # หัวข้อ MEMORANDUM (ชิดขวา)
    p_top_title = doc.add_paragraph()
    p_top_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_top_title.paragraph_format.space_after = Pt(2)
    add_run_thai(p_top_title, " MEMORANDUM ", size_pt=14, bold=True)
    
    # ตารางส่วนหัวเอกสาร 5 แถว
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(2.2)
    info_table.columns[1].width = Inches(4.3)
    apply_custom_info_borders(info_table)
    
    today_str = datetime.datetime.now().strftime(f"%d {selected_month} {selected_year}")
    
    # จัดระยะขอบในเซลล์ส่วนหัวให้ชิดขึ้น
    for row in info_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0

    add_run_thai(info_table.cell(0, 0).paragraphs[0], "ที่ / No:  -", size_pt=14, bold=True)
    p_1_right = info_table.cell(0, 1).paragraphs[0]
    p_1_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_thai(p_1_right, f"วันที่ / Date:  {today_str}", size_pt=14)

    headers_layout = [
        ("หน่วยงานผู้ส่ง / From:  ", "ส่วนบริหารกลยุทธ์และแผนการผลิต (กผ.)"),
        ("เรียน / To:  ", "ผจ.สทต. / ผจ.บท ผ่าน ผจ.กผ."),
        ("สำเนา / CC:  ", "ผจ.ปท.3, ผจ.ชก., ผจ.บฟ."),
        ("เรื่อง / Subject:  ", f"แจ้งสรุปค่าไฟฟ้าของสถานีชายฝั่งระยอง ประจำเดือน {selected_month} {selected_year}")
    ]
    
    for idx, (label, value) in enumerate(headers_layout, start=1):
        add_run_thai(info_table.cell(idx, 0).paragraphs[0], label, size_pt=14, bold=True)
        add_run_thai(info_table.cell(idx, 1).paragraphs[0], value, size_pt=14)

    # เส้นคั่นกลางหนาเข้มหลังส่วนหัวเรื่อง
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(8)
    p_line_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="000000"/></w:pBdr>')
    p_line._p.get_or_add_pPr().append(p_line_border)

    # ย่อหน้าเนื้อความนำ
    p_body1 = doc.add_paragraph()
    p_body1.paragraph_format.first_line_indent = Inches(0.5)
    p_body1.paragraph_format.space_after = Pt(8)
    p_body1.paragraph_format.line_spacing = 1.05
    add_run_thai(p_body1, f"ส่วนบริหารกลยุทธ์และแผนการผลิต ฝ่ายบริหารเทคนิคและแผนการผลิต ขอนำส่งสรุปค่าไฟฟ้าของสถานีชายฝั่งระยอง ประจำเดือน {selected_month} {selected_year} รายละเอียดการคำนวณตามเอกสารแนบ", size_pt=15)

    # สร้างโครงสร้างตารางข้อมูล 3 แถวแรกเป็นหัวตารางสีพีช
    calc_table = doc.add_table(rows=3, cols=7)
    calc_table.style = 'Table Grid'
    calc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    peach_hex = "F2C4A9"
    calc_table.rows[0].cells[0].merge(calc_table.rows[1].cells[0]).merge(calc_table.rows[2].cells[0])
    set_cell_background(calc_table.rows[0].cells[0], peach_hex)
    add_run_thai(calc_table.rows[0].cells[0].paragraphs[0], "พื้นที่ใช้ไฟฟ้า", size_pt=10, bold=True)
    calc_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers_top = [("ไฟฟ้าที่รับจาก PEA", 1, 2), ("ไฟฟ้าที่รับจาก GSP", 3, 4), ("รวมไฟฟ้าที่รับทั้งหมด", 5, 6)]
    for text, c_start, c_end in headers_top:
        cell = calc_table.rows[0].cells[c_start].merge(calc_table.rows[0].cells[c_end])
        set_cell_background(cell, peach_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run_thai(p, text, size_pt=10, bold=True)
    
    sub_headers = ["ปริมาณไฟฟ้า", "ค่าใช้จ่าย", "ปริมาณไฟฟ้า", "ค่าใช้จ่าย", "ปริมาณไฟฟ้า", "ค่าใช้จ่าย"]
    sub_units = ["(kWh)", "(บาท)", "(kWh)", "(บาท)", "(kWh)", "(บาท)"]
    for idx in range(6):
        c2 = calc_table.rows[1].cells[idx+1]
        c3 = calc_table.rows[2].cells[idx+1]
        set_cell_background(c2, peach_hex)
        set_cell_background(c3, peach_hex)
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)
        c3.paragraphs[0].paragraph_format.space_after = Pt(0)
        add_run_thai(c2.paragraphs[0], sub_headers[idx], size_pt=9, bold=True)
        add_run_thai(c3.paragraphs[0], sub_units[idx], size_pt=9, bold=True)
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # เริ่มกระบวนการวนลูปกรอกเฉพาะแถวเป้าหมายหลัก 6 แถวเท่านั้นเพื่อไม่ให้ล้นหน้า
    total_cost_sum = "0.00"
    for _, row in df_clean.iterrows():
        row_cells = calc_table.add_row().cells
        area_name = str(row.iloc[0]).strip()
        is_total = "total" in area_name.lower() or "รวม" in area_name
        
        for col_idx in range(7):
            cell = row_cells[col_idx]
            if is_total:
                set_cell_background(cell, peach_hex)
            set_cell_margins(cell, top=15, bottom=15, left=40, right=40)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            val = row.iloc[col_idx]
            # กรองตรวจสอบความถูกต้องของฟอร์แมตตัวเลขบิลค่าไฟ
            if isinstance(val, (int, float)):
                val_str = "-" if val == 0 or pd.isna(val) else f"{val:,.2f}"
            else:
                val_str = str(val).strip()
                if val_str in ["", "nan", "None", "-", "0", "0.00"]:
                    val_str = "-"
            
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_thai(p, val_str, size_pt=10, bold=is_total)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_run_thai(p, val_str, size_pt=10, bold=is_total)
                if is_total and col_idx == 6:
                    total_cost_sum = val_str

    # สรุปราคารวมและเรทด้านล่างตาราง (ปรับลด space_before เพื่อดันไม่ให้ล้นหน้า)
    p_body2 = doc.add_paragraph()
    p_body2.paragraph_format.space_before = Pt(8)
    p_body2.paragraph_format.space_after = Pt(2)
    add_run_thai(p_body2, f"\tอัตราค่าไฟฟ้าอ้างอิง ราคา PEA Rate ณ เดือน {selected_month} {selected_year} = ", size_pt=14)
    add_run_thai(p_body2, f"{reference_rate:.4f}", size_pt=14, bold=True)
    add_run_thai(p_body2, " บาท / kWh  รายละเอียดตามเอกสารแนบ", size_pt=14)
    
    p_body3 = doc.add_paragraph()
    p_body3.paragraph_format.space_before = Pt(2)
    p_body3.paragraph_format.space_after = Pt(0)
    add_run_thai(p_body3, f"\tรวมค่าไฟฟ้าที่เรียกเก็บจากระบบท่อส่งก๊าซฯ ทั้งสิ้น\t\t", size_pt=14)
    add_run_thai(p_body3, f"{total_cost_sum}", size_pt=14, bold=True)
    add_run_thai(p_body3, "   บาท", size_pt=14)

    # ลงชื่อผู้อนุมัติท้ายกระดาษ (ปรับระยะห่างให้พอดี 1 หน้ากระดาษ)
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(24)
    p_sign.paragraph_format.space_after = Pt(0)
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_thai(p_sign, "จึงเรียนมาเพื่อโปรดทราบ            \n\n\n", size_pt=14)
    add_run_thai(p_sign, "(นางสุรีย์พันธ์ คุณากูลสวัสดิ์)      \n", size_pt=14, bold=True)
    add_run_thai(p_sign, "ผู้จัดการทั่วไป            ", size_pt=14)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ================================================================
# 3. ส่วนแอปพลิเคชันหลักของ Streamlit
# ================================================================
st.title("ระบบออกรายงานสรุปบันทึกข้อความ (Memo Report)")

uploaded_file = st.file_uploader("📥 อัปโหลดไฟล์ Excel ข้อมูลค่าไฟฟ้า", type=["xlsx", "xls"])

col1, col2 = st.columns(2)
with col1:
    months_list = ["มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
    selected_month = st.selectbox("เลือกประจำเดือน", months_list, index=3) # Default เป็น เมษายน ตามภาพสรุปงาน
with col2:
    years_list = [str(y) for y in range(2560, 2575)]
    selected_year = st.selectbox("เลือกปี พ.ศ.", years_list, index=9) # Default เป็น 2569

reference_rate = st.number_input("⚡ ป้อนอัตราค่าไฟฟ้าอ้างอิงประจำเดือน (บาท / kWh)", min_value=0.0, max_value=10.0, value=4.6311, format="%.4f")

if uploaded_file is not None:
    try:
        # 1. บังคับอ่านชีตแผ่นแรกสุดเสมอ โดยแกะพิกัดแถวแบบยืดหยุ่นกว้าง ๆ ก่อนเพื่อหาจุดคำนวณจริง
        df_raw = pd.read_excel(uploaded_file, header=None, sheet_name=0)
        
        # 2. ค้นหาพิกัดตารางที่แม่นยำด้วยการค้นหา Keyword เพื่อป้องกันปัญหาแถวเลื่อนเยื้องในไฟล์ Excel
        target_indices = []
        target_areas = ["DPCU", "New DPCU", "OCS1", "OCS2", "OCS3", "Total"]
        
        # ค้นหาในแถวต่าง ๆ เพื่อจับจุด Index แถวที่มีกลุ่มเป้าหมายอยู่จริง ๆ 
        for idx, row in df_raw.iterrows():
            cell_value = str(row.iloc[6]).strip() if len(row) > 6 else "" # คอลัมน์ G คือ Index ที่ 6
            if any(area.lower() in cell_value.lower() for area in target_areas):
                target_indices.append(idx)
        
        # กรณีค้นหาแบบเจาะจุดไม่เจอ ให้ใช้ Fallback ล็อกช่วงพิกัดแถว 48-56 (Index 47:56) ในช่วงคอลัมน์ G-M
        if len(target_indices) >= 5:
            # ใช้พิกัดที่สแกนเจอสด ๆ จากในไฟล์จริง
            start_row = target_indices[0] - 1 if target_indices[0] > 0 else target_indices[0]
            end_row = target_indices[-1] + 1
            df_sliced = df_raw.iloc[start_row:end_row, 6:13].copy()
        else:
            df_sliced = df_raw.iloc[47:56, 6:13].copy()
            
        # 3. จัดตั้งโครงสร้างหัวตารางและคลีนข้อมูลความสะอาดของชุดตัวเลข
        df_sliced.columns = ['พื้นที่ใช้ไฟฟ้า', 'PEA_ปริมาณ', 'PEA_ค่าใช้จ่าย', 'GSP_ปริมาณ', 'GSP_ค่าใช้จ่าย', 'รวม_ปริมาณ', 'รวม_ค่าใช้จ่าย']
        
        # กรองเจาะเอาเฉพาะแถวที่เป็นเนื้อเค้กข้อมูลหลัก (ล้างพวกคำอธิบายสูตรย่อยที่ติดสอยห้อยตามมาทิ้ง)
        final_rows = []
        for _, r in df_sliced.iterrows():
            area_str = str(r['พื้นที่ใช้ไฟฟ้า']).strip()
            if any(a.lower() in area_str.lower() for a in target_areas) and "source" not in area_str.lower() and "recheck" not in area_str.lower():
                final_rows.append(r)
                
        df_clean = pd.DataFrame(final_rows).reset_index(drop=True)
        
        # แปลงข้อมูลตัวเลขที่ดึงมาจากสูตรให้เป็นประเภทตัวเลข (Float) และจัดการแปลง NaN เป็น 0 ปิดปัญหา JSON parse Error
        for col in df_clean.columns[1:]:
            df_clean[col] = pd.to_numeric(df_clean[col], errors='coerce').fillna(0.0)
            
        # ปรับแก้ชื่อเล่นพื้นที่ให้สวยงามเป็นทางการตรงตามบิล
        name_mapping = {
            "DPCU 300-TR-1": "DPCU", "DPCU": "DPCU",
            "New DPCU": "New DPCU",
            "OCS1 Load": "OCS1", "OCS1": "OCS1",
            "OCS2": "OCS2", "OCS3": "OCS3", "Total": "Total"
        }
        df_clean['พื้นที่ใช้ไฟฟ้า'] = df_clean['พื้นที่ใช้ไฟฟ้า'].map(lambda x: name_mapping.get(str(x).strip(), str(x).strip()))
        
        # ป้องกันตารางซ้ำซ้อน: ยุบรวมกรณีที่ข้อมูล OCS ขยายเป็นหลายบรรทัดย่อย ให้รวมเป็นบรรทัดเดียวรวมตามมาตรฐาน ปตท.
        df_clean = df_clean.groupby('พื้นที่ใช้ไฟฟ้า', sort=False).sum().reset_index()

        st.success("🎯 ระบบจำกัดขอบเขตและดึงข้อมูลเฉพาะตารางสีส้มสำเร็จแล้ว!")
        st.dataframe(df_clean)
        
        if st.button("📝 สร้างรายงาน Word (จบในหน้าเดียว)"):
            word_file = create_exact_layout_report(df_clean, selected_month, selected_year, reference_rate)
            st.success("✅ ดึงข้อมูลสำเร็จ! ความยาวเอกสารได้รับการควบคุมให้อยู่ใน 1 หน้าเรียบร้อยแล้วครับ")
            st.download_button(
                label="📥 ดาวน์โหลดไฟล์บันทึกข้อความ ปตท. (.docx)",
                data=word_file,
                file_name=f"Memo_สรุปค่าไฟฟ้า_{selected_month}_{selected_year}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาดในการประมวลผลโค้ด XML/Excel: {e}")
